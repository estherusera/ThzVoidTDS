"""
evaluation_metrics_doc.py — generate evaluation_metrics.docx
Defines IoU, mIoU, Dice, precision, recall, pixel accuracy for the void-
segmentation thesis, with formulas, relationships, caveats, and references.
Run: thesis_env/bin/python evaluation_metrics_doc.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path("evaluation_metrics.docx")
doc = Document()
doc.styles["Normal"].font.name = "Times New Roman"
doc.styles["Normal"].font.size = Pt(11)


def H1(t): return doc.add_heading(t, level=1)
def H2(t): return doc.add_heading(t, level=2)
def P(t): return doc.add_paragraph(t)


def EQ(t):
    """Centered italic equation line."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t); r.italic = True; r.font.size = Pt(12)
    return p


def BULLET(t): return doc.add_paragraph(t, style="List Bullet")


doc.add_heading("Evaluation Metrics for Void Segmentation", level=0)
P("All segmentation metrics used in this work are derived from the per-pixel "
  "(per-voxel, for the 3-D volumes) confusion matrix of the binary problem "
  "void vs. background. For each pixel the prediction is compared with the "
  "ground-truth label, yielding four counts: true positives (TP, predicted "
  "void and truly void), false positives (FP, predicted void but background), "
  "false negatives (FN, predicted background but truly void), and true "
  "negatives (TN, predicted background and truly background). Every metric "
  "below is a function of these four counts.")

H1("Confusion matrix")
tbl = doc.add_table(rows=3, cols=3)
tbl.style = "Light Grid Accent 1"
cells = [["", "Ground truth: void", "Ground truth: background"],
         ["Predicted: void", "TP (true positive)", "FP (false positive)"],
         ["Predicted: background", "FN (false negative)", "TN (true negative)"]]
for i, row in enumerate(cells):
    for j, txt in enumerate(row):
        c = tbl.rows[i].cells[j]; c.text = txt
        if i == 0 or j == 0:
            for r in c.paragraphs[0].runs: r.bold = True

H1("Metric definitions")

H2("Precision")
EQ("Precision = TP / (TP + FP)")
P("Of the pixels the model labelled as void, the fraction that are truly void. "
  "Low precision indicates over-segmentation (false alarms).")

H2("Recall (sensitivity, true-positive rate)")
EQ("Recall = TP / (TP + FN)")
P("Of the pixels that are truly void, the fraction the model recovered. Low "
  "recall indicates missed voids. Precision and recall trade off against each "
  "other and should be reported together.")

H2("Pixel accuracy")
EQ("PA = (TP + TN) / (TP + TN + FP + FN)")
P("The fraction of all pixels classified correctly. Because voids occupy only a "
  "small fraction of each scan, this metric is misleading under such class "
  "imbalance: a model predicting background everywhere attains very high "
  "accuracy while detecting nothing. It is reported for completeness but is not "
  "relied upon in isolation.")

H2("Intersection-over-Union (IoU, Jaccard index)")
EQ("IoU = TP / (TP + FP + FN) = |A ∩ B| / |A ∪ B|")
P("The overlap between the predicted void region A and the ground-truth void "
  "region B, normalised by their union. It ranges from 0 (no overlap) to 1 "
  "(perfect overlap). Unlike pixel accuracy, IoU excludes the dominant true-"
  "negative background and therefore genuinely measures void localisation.")

H2("Dice coefficient (Sørensen–Dice; equivalent to the F1 score)")
EQ("Dice = 2·TP / (2·TP + FP + FN) = 2|A ∩ B| / (|A| + |B|)")
EQ("Dice = 2·Precision·Recall / (Precision + Recall)")
P("Twice the intersection divided by the sum of the two sets. As the third form "
  "shows, the Dice coefficient is the harmonic mean of precision and recall and "
  "is therefore identical to the F1 score; it is high only when both precision "
  "and recall are high.")

H2("Mean IoU (mIoU)")
EQ("mIoU = (1/N) · Σ IoUᵢ ,  i = 1 … N")
P("The IoU averaged over a set of N items. In a two-class problem this is the "
  "mean of the void IoU and the background IoU; in a multi-sample report it is "
  "the mean of the per-sample void IoU. The averaging convention should be "
  "stated explicitly, as both appear in the literature; in this work mIoU "
  "denotes the per-sample mean of the void-class IoU.")

H1("Relationships between the metrics")
P("Dice and IoU are monotonically related, so they rank models identically:")
EQ("Dice = 2·IoU / (1 + IoU) ,   IoU = Dice / (2 − Dice)")
BULLET("Consequently Dice ≥ IoU for any imperfect prediction; Dice is the "
       "more lenient-looking value. Both are commonly reported, but they carry "
       "the same ranking information.")
BULLET("Dice and IoU involve only the positive (void) class, making them the "
       "appropriate choice under heavy class imbalance, whereas pixel accuracy "
       "is dominated by true negatives.")

H1("Application note")
P("Because the majority of depth slices contain no void, per-slice Dice and IoU "
  "are trivially close to 1 on empty slices and inflate the mean. Metrics in "
  "this work are therefore reported void-conditionally (computed over slices or "
  "samples that contain at least one void), and precision and recall are "
  "reported separately, since the detectors are characteristically high-"
  "precision / low-recall — a property that a single Dice value obscures.")

H1("References")
refs = [
 "Jaccard, P. (1912). The distribution of the flora in the alpine zone. New "
 "Phytologist, 11(2), 37–50.",
 "Dice, L. R. (1945). Measures of the amount of ecologic association between "
 "species. Ecology, 26(3), 297–302.",
 "Sørensen, T. (1948). A method of establishing groups of equal amplitude "
 "in plant sociology based on similarity of species content. Biologiske "
 "Skrifter, 5, 1–34.",
 "van Rijsbergen, C. J. (1979). Information Retrieval (2nd ed.). Butterworths.",
 "Everingham, M., Van Gool, L., Williams, C. K. I., Winn, J., & Zisserman, A. "
 "(2010). The PASCAL Visual Object Classes (VOC) Challenge. International "
 "Journal of Computer Vision, 88(2), 303–338.",
 "Long, J., Shelhamer, E., & Darrell, T. (2015). Fully Convolutional Networks "
 "for Semantic Segmentation. CVPR, 3431–3440.",
 "Taha, A. A., & Hanbury, A. (2015). Metrics for evaluating 3D medical image "
 "segmentation: analysis, selection, and tool. BMC Medical Imaging, 15, 29.",
 "Milletari, F., Navab, N., & Ahmadi, S.-A. (2016). V-Net: Fully Convolutional "
 "Neural Networks for Volumetric Medical Image Segmentation. 3DV, 565–571.",
 "Csurka, G., Larlus, D., & Perronnin, F. (2013). What is a good evaluation "
 "measure for semantic segmentation? BMVC.",
]
for r in refs:
    doc.add_paragraph(r, style="List Number")

doc.save(OUT)
print(f"wrote {OUT}")
