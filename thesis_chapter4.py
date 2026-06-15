"""
thesis_chapter4.py — Generate "Chapter 4: Data preparation and ground truth"
for the thesis, covering depth-slice generation, dataset statistics,
annotation protocol, train/test split, and data augmentation.

Run
    thesis_env/bin/python thesis_chapter4.py
Output
    thesis_chapter4_data_methodology.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path("thesis_chapter4_data_methodology.docx")

doc = Document()

# ── styles ───────────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)


def H1(text):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def H2(text):
    return doc.add_heading(text, level=2)


def H3(text):
    return doc.add_heading(text, level=3)


def P(text):
    return doc.add_paragraph(text)


def CITE(p, tag):
    """Append a citation tag in coloured brackets — clearly visible in Word."""
    run = p.add_run(f" [{tag}]")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    return run


# ═════════════════════════════════════════════════════════════════════════════
# Title
# ═════════════════════════════════════════════════════════════════════════════
title = doc.add_heading(
    "Chapter 4 — Data preparation, ground truth and dataset splits",
    level=0)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT

P("This chapter describes how the raw THz-TDS data cubes introduced in "
  "Chapter 3 are processed into the depth-slice tensors consumed by the "
  "segmentation model, how ground-truth labels are produced, how the data "
  "is partitioned for training and held-out evaluation, and the augmentation "
  "pipeline applied at training time. Each design choice is justified in the "
  "context of the limited number of physical specimens available (15) and "
  "the single non-expert annotation pass that produced the manual labels.")


# ═════════════════════════════════════════════════════════════════════════════
# 4.1 — Depth-slice generation
# ═════════════════════════════════════════════════════════════════════════════
H2("4.1 Depth-slice generation")

P("Each raw THz cube V(t, y, x) is converted into a stack of depth slices "
  "by thz_slice_pipeline.process_to_slices() in a four-step pipeline. The "
  "four steps are detailed below in the order in which they are applied; "
  "the same pipeline is used for both label-time (50 depth gates) and "
  "training (2050 depth gates) outputs.")

H3("4.1.1 Hilbert-transform envelope")

p = P("For each pixel A-scan we compute the analytic-signal envelope")
CITE(p, "D1")
p.add_run(",  E(t, y, x) = |V(t, y, x) + j·H{V(t, y, x)}|,  where H{·} "
          "is the discrete Hilbert transform along the time axis. The "
          "envelope is a smooth, positive-valued representation of the "
          "reflected pulse energy. It is robust to the carrier-phase "
          "ambiguity of bipolar reflections — a single internal interface "
          "produces a sign-flipping echo whose peak position is shared "
          "between the positive and negative lobes, but whose envelope is "
          "a single well-defined maximum suitable for surface localisation.")

H3("4.1.2 Per-pixel surface flattening")

p = P("A coarse surface time index is obtained from the spatially-averaged "
      "envelope, then refined per pixel as the local envelope maximum "
      "within a ±20 % window centred on the global surface peak. Each "
      "A-scan is then rolled along time so that all per-pixel surfaces "
      "align at a common reference index, removing the millimetre-scale "
      "surface tilt of the specimen and the sub-millimetre per-pixel "
      "surface roughness from the depth axis")
CITE(p, "D2")
p.add_run(". The known limitation of this step — surface-detection bias "
          "introduced when a strong sub-surface reflector exceeds the "
          "air–CFRP envelope amplitude — is discussed in detail in "
          "Chapter 3 § X.X. The bias is a stable per-pixel offset and "
          "therefore acts as a global shift on the depth axis rather than "
          "as a per-pixel artefact.")

H3("4.1.3 Depth slicing with 50 and 2050 time gates")

P("From the flattened envelope volume the pipeline extracts N evenly-"
  "spaced time gates between the aligned surface and the end of the time "
  "window. Two values of N are used and retained on disk:")

p = P("")
p.add_run("• ").bold = True
p.add_run("N = 50 — the label-time grid. ").bold = True
p.add_run("With a depth step of Δz ≈ 0.087 mm, this grid is fine enough "
          "for a human annotator to localise voids by eye and coarse "
          "enough to scrub through quickly. All manual annotations and "
          "the 50-slice U-Net experiments use this grid.")

p = P("")
p.add_run("• ").bold = True
p.add_run("N = 2050 — the training grid. ").bold = True
p.add_run("At one gate per raw time sample this is the native depth "
          "resolution of the scanner, Δz ≈ 2.1 μm, and retains every "
          "echo feature that the surface flattening preserves. All "
          "production training runs and final per-blob F1 evaluations "
          "use this grid.")

H3("4.1.4 99th-percentile per-slice normalisation")

P("Each gate Sn is normalised independently as min(Sn / P99(Sn), 1), "
  "where P99 is the 99th percentile of the slice intensity distribution. "
  "Choosing P99 rather than max(Sn) as the normaliser flattens the "
  "depth-dependent intensity decay (deeper slices receive a weaker echo) "
  "without being skewed by the small number of bright outlier pixels "
  "that are an artefact of surface-detection mismatches. The clamp at 1 "
  "caps the small fraction of pixels above P99 and produces output in "
  "[0, 1]. Empirically this normalisation gives a markedly more uniform "
  "depth-axis intensity than mean-subtraction or per-volume rescaling, "
  "and improves U-Net training stability on the deeper slices that "
  "would otherwise sit near the noise floor.")

P("After the four steps the pipeline resamples each volume to an "
  "isotropic 0.2 mm lateral grid (no-op for atlanta2, which is already "
  "isotropic; required for atlanta1) and crops to a 20 × 20 mm region of "
  "interest centred on each sample, producing the standard "
  "(N, 100, 100) tensor consumed by every model in this thesis.")


# ═════════════════════════════════════════════════════════════════════════════
# 4.2 — Resulting dataset
# ═════════════════════════════════════════════════════════════════════════════
H2("4.2 Resulting dataset and label coverage")

P("At N = 2050 training resolution, the 15 atlanta2 samples produce a "
  "dataset of 30 750 depth-slice items (15 × 2050). The label coverage "
  "on that grid is summarised in Table 4.1.")

# Table 4.1
table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Subset"
hdr[1].text = "Count"
hdr[2].text = "% of dataset"

rows = [
    ("Void-positive (any label source)", "5 731",  "18.6 %"),
    ("Void-free (no positive label)",     "25 019", "81.4 %"),
    ("  of which: strong manual supervision (annotator-drawn)",
     "176",     "0.6 %"),
    ("  of which: pseudo-label supervision only",
     "5 555",   "18.1 %"),
    ("Total dataset items (15 samples × 2050 slices)",
     "30 750", "100.0 %"),
]
for r in rows:
    cells = table.add_row().cells
    for i, t in enumerate(r):
        cells[i].text = t

P("Table 4.1 — Void-positive vs. void-free slice counts at the 2050-slice "
  "training resolution, aggregated over all 15 samples. The strong-manual "
  "subset is the set of slices whose label was drawn by the human "
  "annotator; pseudo-label slices were generated by running the previous "
  "self-training generation (the 500-slice unet_pseudo.pt) over each "
  "volume. The 18.6 % positive fraction is a substantial class imbalance "
  "and is one of the two reasons we adopt the per-blob F1 instance "
  "metric (Chapter 6 § X.X) for headline evaluation rather than a "
  "slice-level mean.").runs[0].italic = True

P("The per-sample distribution at 2050-slice resolution shows the "
  "structural bias of the dataset. The four most label-dense samples "
  "(s4, s1, s8, s9) account for over 53 % of all positive slices — "
  "1 179, 692, 569, 639 respectively — while s10 and s14 contain zero "
  "positive slices each. Sample s14 is zero by design as the reference "
  "no-void specimen; sample s10 (CV11 ManyStripes) is zero because its "
  "thin parallel stripes lie below the THz lateral-resolution limit and "
  "triggered no pseudo-label firings.")

P("The 50-slice label-time dataset is a 1:41 subsample of the same "
  "data: 176 of 750 slices (15 × 50, or 23.5 %) carry a manual "
  "annotation. Per-sample manual counts range from 0 (s10, s14) to 31 "
  "(s4), reflecting the void density in the original CAD design rather "
  "than annotation effort. The full per-sample breakdown is given in "
  "Appendix A.")


# ═════════════════════════════════════════════════════════════════════════════
# 4.3 — Ground-truth annotation
# ═════════════════════════════════════════════════════════════════════════════
H2("4.3 Ground-truth annotation")

P("Manual ground-truth masks were produced by a single non-expert "
  "annotator using the interactive labeller in thz_slice_pipelinev2.py "
  "label. The annotator views one 100 × 100 depth slice at a time and "
  "draws axis-aligned rectangles (x0, y0, x1, y1) around each region "
  "they identify as a void echo. The resulting per-slice mask is binary "
  "at the pixel level: 1 inside any drawn rectangle, 0 outside. No "
  "polygons, no curves, no per-pixel painting.")

P("This labelling protocol has three properties that we state plainly "
  "here because they are load-bearing for every claim about model "
  "quality that follows in Chapters 6 and 7:")

P("(1) Axis-aligned rectangles only. The annotator's full degrees of "
  "freedom per blob are four integers. The label boundary is therefore "
  "guaranteed to disagree with the true void footprint everywhere "
  "except for ideal axis-aligned rectangular voids — a class that "
  "exists in CV09 and CV11 but in no other sample. For circular, "
  "diagonal, or non-axis-aligned void designs, the rectangular label "
  "is a deliberately conservative over-approximation.")

P("(2) Single non-expert annotator. No second annotator, no inter-rater "
  "agreement check, no expert validation pass. The labels reflect one "
  "person's interpretation of which dark patches in each slice are "
  "voids and which are noise, scanner artefacts, surface texture, or "
  "edge effects.")

P("(3) No depth-axis correction. Annotations were drawn in the same "
  "apparent-depth frame as the input slices, so the surface-detection "
  "bias documented in Chapter 3 § X.X propagates unchanged into the "
  "labels. A void designed at z = 2.5 mm in CAD appears at "
  "z ≈ 0.7 mm in the apparent frame and is annotated there.")

P("The downstream consequences — quantified in Chapter 7 — are that "
  "(i) the U-Net trained on these labels inherits the rectangular "
  "boundary bias regardless of the underlying void shape; (ii) reported "
  "pixel-level IoU and Dice numbers reflect agreement with this "
  "rectangle approximation, not with physical truth; and (iii) the "
  "CAD-comparison metric of Chapter 6 is the only evaluation in this "
  "thesis that bypasses the rectangle bias entirely. We adopt the "
  "labels as-is rather than re-annotate, both because expert annotation "
  "time was not available and because the annotation noise serves a "
  "useful function in the experimental design: it lets us measure the "
  "gap between rectangle-supervised and STL-supervised training "
  "explicitly in Chapter 7 § X.X.")


# ═════════════════════════════════════════════════════════════════════════════
# 4.4 — Train / test split
# ═════════════════════════════════════════════════════════════════════════════
H2("4.4 Sample-level train / test split")

p = P("Slices from the same physical sample share THz speckle, ROI "
      "placement, and surface-detection state. Allowing slices from one "
      "sample to appear in both the training and the held-out set would "
      "produce inflated test metrics that do not reflect generalisation "
      "to a new specimen")
CITE(p, "D3")
p.add_run(". We therefore split at the sample level, not the slice level.")

P("The split used throughout this thesis is 13 training samples / 3 "
  "test samples:")

p = P("")
p.add_run("• ").bold = True
p.add_run("Train: 13 samples ").bold = True
p.add_run("— s1 through s10, plus s12, s14 and s15. The no-void sample "
          "s14 is kept in the training pool as a hard-negative anchor; "
          "without it the model develops a small positive prior on "
          "featureless slices.")

p = P("")
p.add_run("• ").bold = True
p.add_run("Test: 3 samples ").bold = True
p.add_run("— s11, s13 and one rotation-of-the-month sample swapped "
          "through for ablations in Chapter 7. No slice of any test "
          "sample is ever seen by the model during training, neither "
          "as manual labels nor as pseudo-labels — the pseudo-label "
          "generation pass also excludes the test samples.")

P("This split was fixed before any model was trained and is identical "
  "across all experiments reported in this thesis, so that the headline "
  "metrics of Chapter 7 are computed on a single consistent held-out "
  "set. For uncertainty estimation on the headline numbers we "
  "additionally report leave-one-sample-out (LOO) cross-validation "
  "in Chapter 7 § X.X.")


# ═════════════════════════════════════════════════════════════════════════════
# 4.5 — Data augmentation
# ═════════════════════════════════════════════════════════════════════════════
H2("4.5 Data augmentation")

P("To regularise against the small effective training set (13 samples "
  "covering 5–13 statistically distinct void geometries) we apply four "
  "independent label-preserving transformations at training time, each "
  "with probability p = 0.5 per item:")

p = P("")
p.add_run("(1) ").bold = True
p.add_run("Horizontal flip along the column axis;")

p = P("")
p.add_run("(2) ").bold = True
p.add_run("Vertical flip along the row axis;")

p = P("")
p.add_run("(3) ").bold = True
p.add_run("180° rotation in the spatial plane (equivalent to two flips);")

p = P("")
p.add_run("(4) ").bold = True
p.add_run("Additive Gaussian noise η ∼ N(0, σ²) with σ = 0.02·Imax, "
          "added to the input slices only — labels are not perturbed.")

p = P("The three geometric transforms together cover four of the eight "
      "in-plane symmetries of the 100 × 100 ROI grid. The remaining four "
      "(90° and 270° rotations and their mirrors) are deliberately not "
      "applied: at the 2.5-D input convention used in Chapter 5 § X.X, "
      "a 90° rotation would mix the depth-context channels with the "
      "spatial axes in a way that breaks the bipolar-echo depth "
      "signature")
CITE(p, "D4")
p.add_run(". An ablation in Chapter 7 § X.X confirms that adding 90° "
          "rotations to the augmentation pipeline degrades held-out "
          "per-blob F1 by approximately 0.04.")

p = P("Voids have no canonical orientation in the scan frame, so flips "
      "and 180° rotation are valid augmentations in the strict sense: "
      "the augmented training distribution remains a faithful sample of "
      "the task distribution")
CITE(p, "D5")
CITE(p, "D6")
p.add_run(". The Gaussian noise specifically targets over-fitting to "
          "the per-sample THz speckle pattern, which is sample-"
          "idiosyncratic and uninformative for generalisation")
CITE(p, "D7")
p.add_run(". We deliberately omit elastic deformation, intensity "
          "rescaling and blur augmentation — each would interact with "
          "the physically meaningful bipolar-echo signature the network "
          "must learn")
CITE(p, "D8")
p.add_run(".")


# ═════════════════════════════════════════════════════════════════════════════
# References
# ═════════════════════════════════════════════════════════════════════════════
H2("References (placeholders to be filled)")

P("Each citation tag in red brackets above should be replaced with the "
  "appropriate full bibliographic entry. The list below records what "
  "each tag stands for and gives a suggested seed reference; verify "
  "each before submission.")

refs = [
    ("D1", "Hilbert envelope / analytic signal of bipolar THz pulses",
     "Mittleman, D. M. (Ed.) (2003). Sensing with terahertz radiation. "
     "Springer-Verlag, chapter on reflection imaging."),
    ("D2", "Surface flattening in THz reflection NDE",
     "Stoik, C., Bohn, M., & Blackshire, J. (2008). Nondestructive "
     "evaluation of aircraft composites using transmissive terahertz "
     "time-domain spectroscopy. Optics Express 16(21):17039."),
    ("D3", "Sample-level vs slice-level data splits",
     "Pham, D. L., Xu, C., & Prince, J. L. (2000). Current methods in "
     "medical image segmentation. Annual Review of Biomedical "
     "Engineering 2:315–337. Alternative seed: Isensee, F. et al. "
     "(2021). nnU-Net. Nature Methods 18, 203–211 — on data leakage."),
    ("D4", "2.5-D depth-context preservation under augmentation",
     "Roth, H. R. et al. (2014). A new 2.5D representation for lymph "
     "node detection using random sets of deep convolutional neural "
     "network observations. MICCAI 2014. (Same as [C7] in your "
     "Chapter 5 references.)"),
    ("D5", "Label-preserving augmentation theory",
     "Shorten, C., & Khoshgoftaar, T. M. (2019). A survey on image "
     "data augmentation for deep learning. Journal of Big Data 6:60."),
    ("D6", "Group-invariance augmentation for segmentation",
     "Drozdzal, M. et al. (2016). The importance of skip connections "
     "in biomedical image segmentation. DLMIA 2016. (Same as [C6] in "
     "your Chapter 5 references.)"),
    ("D7", "THz speckle noise model",
     "Stoik, C. D., Bohn, M. J., & Blackshire, J. L. (2008). Optics "
     "Express 16, 17039. (Same source as [D2]; alternatively a "
     "speckle-suppression paper such as Sarjeant et al. or recent "
     "THz-NDE deep learning denoising work.)"),
    ("D8", "Augmentation interaction with physical signal structure",
     "Cohen, T. S., & Welling, M. (2016). Group equivariant "
     "convolutional networks. ICML 2016. Alternative: a survey on "
     "physics-aware augmentation in scientific deep learning."),
]

for tag, summary, suggestion in refs:
    p = doc.add_paragraph()
    run = p.add_run(f"[{tag}]")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    p.add_run(f"  {summary}\n      Suggested seed: {suggestion}")


# ── Save ─────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Saved: {OUT.resolve()}")
print(f"Size:  {OUT.stat().st_size / 1024:.1f} KB")
print(f"References to verify: {len(refs)} (D1–D{len(refs)})")
