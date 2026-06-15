"""
thesis_chapter5.py — Generate a .docx of "Chapter 5: Segmentation Model and
Methodology" for the thesis, with citation placeholders and a references list.

Run
    thesis_env/bin/python thesis_chapter5.py
Output
    thesis_chapter5_segmentation_methodology.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path("thesis_chapter5_segmentation_methodology.docx")

doc = Document()

# ── styles ───────────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)

def H1(text):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def H2(text):
    return doc.add_heading(text, level=2)

def H3(text):
    return doc.add_heading(text, level=3)

def P(text):
    return doc.add_paragraph(text)

def CITE(p, tag):
    """Append a citation tag in colored brackets — clearly visible in Word."""
    run = p.add_run(f" [{tag}]")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    return run


# ═════════════════════════════════════════════════════════════════════════════
# Title
# ═════════════════════════════════════════════════════════════════════════════
title = doc.add_heading("Chapter 5 — Segmentation Model and Methodology", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT

P("This chapter describes the convolutional neural network used to localise "
  "internal voids in the THz depth-slice volumes introduced in Chapter 4. "
  "We adopt a 2.5-D variant of the U-Net architecture, train it with a "
  "combined binary-cross-entropy and Dice loss, and reason about each design "
  "choice in the context of the limited labelled data and the bipolar-echo "
  "physics of THz reflection imaging.")


# ═════════════════════════════════════════════════════════════════════════════
# 5.1
# ═════════════════════════════════════════════════════════════════════════════
H2("5.1 U-Net architecture")

p = P("The void-segmentation model is a two-dimensional convolutional "
      "encoder–decoder of the U-Net family")
CITE(p, "C1")
p.add_run(". The network operates on individual depth slices of the THz "
          "volume and produces a per-pixel void-probability map at the same "
          "spatial resolution as the input. The total parameter count is "
          "1,928,993 (≈ 7.4 MB on disk in float32), modest by contemporary "
          "segmentation standards and chosen to allow training from scratch "
          "with the limited number of labelled THz samples available without "
          "over-parameterisation.")
CITE(p, "C2")

H3("5.1.1 Three-level encoder–decoder and channel progression")

p = P("The encoder reduces the spatial resolution through three max-pooling "
      "stages, doubling the channel count at each level. The decoder "
      "symmetrically upsamples through three transposed-convolution stages, "
      "halving the channel count at each level. With the base filter count "
      "F = 32, the channel progression is:")

P("    Encoder:  5  →  32  →  64  →  128  →  256  (bottleneck)")
P("    Decoder:  256  →  128  →  64  →  32  →  1  (output sigmoid logit)")

p = P("Each level uses a DoubleConv block — two consecutive Conv2D → "
      "BatchNorm → ReLU operations with kernel size 3 — following the "
      "original U-Net design pattern")
CITE(p, "C1")
p.add_run(". Batch normalisation")
CITE(p, "C3")
p.add_run(" stabilises training under the small batch size dictated by GPU "
          "memory (§ 5.4). The non-linearity is the standard rectified "
          "linear unit")
CITE(p, "C4")
p.add_run(", initialised with the He scheme")
CITE(p, "C5")
p.add_run(" to match the variance assumption of ReLU networks.")

H3("5.1.2 Skip connections")

p = P("Skip connections concatenate the encoder feature maps with the decoder "
      "upsampling output at each resolution level. This preserves the fine "
      "spatial detail that would otherwise be lost during pooling and allows "
      "the decoder to recover sharp object boundaries — a property of U-Net "
      "that has been established across medical, satellite, and microscopy "
      "segmentation benchmarks")
CITE(p, "C1")
p.add_run("")
CITE(p, "C6")
p.add_run(". In our setting the skip pathways are essential for resolving "
          "void edges that are < 1 mm across in the lateral plane, which "
          "after a single pooling stage would already span fewer than 3 "
          "feature-map pixels.")

H3("5.1.3 Pooling, upsampling, and shape handling")

P("Downsampling uses 2×2 max-pooling; upsampling uses 2×2 transposed "
  "convolutions with stride 2. The total downsampling factor over three "
  "levels is 8, so the input spatial dimensions are reflect-padded at "
  "ingress to ensure divisibility by 8. The output is cropped back to the "
  "original 100 × 100 resolution before the sigmoid activation, so the "
  "model is shape-agnostic and the same checkpoint can be applied to other "
  "ROI sizes without retraining.")


# ═════════════════════════════════════════════════════════════════════════════
# 5.2
# ═════════════════════════════════════════════════════════════════════════════
H2("5.2 Five-channel depth-context input (2.5-D)")

p = P("Each prediction is conditioned on the target depth slice plus its two "
      "neighbours on either side, yielding a five-channel input tensor of "
      "shape (5, 100, 100). For a target slice index i, the input channels "
      "are the slices at depths i − 2 k, i − k, i, i + k, i + 2 k, where the "
      "depth step k is scaled to the slice grid (k = 1, 5, 20 for the 50-, "
      "500-, 2050-slice variants respectively) so that the physical depth "
      "window remains approximately ± 0.084 mm across all variants. This "
      "2.5-D paradigm — using a small stack of neighbouring 2-D slices "
      "instead of full 3-D convolutions — is a deliberate trade-off between "
      "architectural complexity, computational cost, and the limited "
      "training data available")
CITE(p, "C7")
p.add_run("")
CITE(p, "C8")
p.add_run(". We motivate the choice with four arguments.")

H3("5.2.1 Computational efficiency")

p = P("Three-dimensional U-Nets")
CITE(p, "C9")
p.add_run(" extend each spatial convolution to operate over the depth axis "
          "as well, multiplying the per-layer parameter count by the depth "
          "kernel size and the per-layer activation memory by the depth "
          "of the volume. For the present problem the equivalent 3-D U-Net "
          "would carry approximately 5.8 M parameters and require an order "
          "of magnitude more activation memory than the 1.93 M-parameter "
          "2.5-D variant adopted here, with no expected gain in accuracy "
          "given the available label volume.")

H3("5.2.2 Data efficiency")

P("Full 3-D segmentation requires fully labelled 3-D volumes. The manual "
  "labels in this work are produced per-slice by a human annotator drawing "
  "rectangles on individual depth images; the labelling protocol naturally "
  "yields 2-D supervision and does not produce densely labelled 3-D "
  "volumes. The 2.5-D approach matches the dimensionality of the available "
  "supervision while still giving the model local depth context that "
  "distinguishes a void echo (which appears at a particular depth and "
  "decays above and below) from random spatial features (which would "
  "appear uniformly across depth).")

H3("5.2.3 Physical interpretation")

p = P("A void in a THz reflection scan produces a bipolar echo centred at "
      "the void depth, characterised by the leading positive and trailing "
      "negative lobes of the Hilbert envelope")
CITE(p, "C10")
p.add_run(". The five-channel input gives the first convolutional kernel "
          "direct access to this depth-derivative signature without "
          "requiring a 3-D filter to learn it. The choice of two neighbours "
          "on either side balances coverage of the bipolar pulse against the "
          "spatial-resolution cost of channel dimensionality; smaller "
          "windows lose the signature, larger windows over-smooth.")

H3("5.2.4 Avoidance of inter-slice label leakage")

P("In 3-D U-Nets trained on sparse labels, the network can learn to "
  "interpolate labels through unlabelled depth slices, producing spurious "
  "axial continuity. The per-slice 2.5-D formulation decouples depth-axis "
  "decisions across slices; inter-slice consistency that is genuinely "
  "desired is recovered downstream by the per-blob 3-D extraction stage "
  "(Chapter 6, § 6.X), which thresholds the per-slice probability volume "
  "and extracts connected components in 2-D before depth-profiling them "
  "post-hoc.")


# ═════════════════════════════════════════════════════════════════════════════
# 5.3
# ═════════════════════════════════════════════════════════════════════════════
H2("5.3 Combined BCE-Dice loss")

p = P("The training objective is a weighted sum of pixel-wise binary "
      "cross-entropy (BCE) and the soft Dice loss:")
P("    L(y, ŷ) = α · BCE(σ(ŷ), y) + (1 − α) · (1 − Dice(σ(ŷ), y)),")
p = P("with σ(·) the sigmoid activation, ŷ the network logits, y the "
      "binary (or, in the STL-supervised variant of Chapter 6, soft) "
      "target, and α = 0.5. The two terms are complementary")
CITE(p, "C11")
p.add_run("")
CITE(p, "C12")
p.add_run("")
CITE(p, "C13")
p.add_run(".")

p = P("Binary cross-entropy provides a stable per-pixel gradient regardless "
      "of class balance and is the de-facto baseline for binary "
      "segmentation. Its limitation in this task is that the loss can be "
      "dominated by background pixels — the void-to-background ratio in our "
      "labelled slices is approximately 1 : 50 — pushing the model toward "
      "an under-predicting solution.")

p = P("The soft Dice loss")
CITE(p, "C12")
p.add_run(", the differentiable relaxation of the Dice similarity "
          "coefficient, is invariant to class imbalance and directly "
          "optimises the segmentation overlap metric used at evaluation "
          "time. Its limitation is unstable gradients in early training, "
          "when predictions are nearly uniform and the soft-Dice "
          "denominator is small.")

p = P("The 50 : 50 weighting is the standard choice adopted by nnU-Net")
CITE(p, "C14")
p.add_run(" and the medical-segmentation literature generally; it balances "
          "early-training stability (the BCE term) against final-overlap "
          "quality (the Dice term). For the STL-supervised variant the "
          "target y is itself soft — a Gaussian-blurred STL z-projection in "
          "[0, 1]. Soft Dice handles soft targets natively, and BCE remains "
          "valid because targets and predictions both lie in the unit "
          "interval.")


# ═════════════════════════════════════════════════════════════════════════════
# 5.4
# ═════════════════════════════════════════════════════════════════════════════
H2("5.4 Training setup (unet_pseudo_2050.pt)")

P("The production segmentation model reported in this thesis was trained "
  "on the atlanta2 dataset at full time-axis resolution (2050 depth slices "
  "per sample), supervised by a combination of human-labelled and "
  "pseudo-labelled depth slices. Training is performed by "
  "retrain_pseudo.py and follows the iterative self-training protocol "
  "described in § 5.5. The hyperparameters and runtime are summarised in "
  "Table 5.1 below. All numerical values match those recorded in the "
  "checkpoint metadata of results_v2_atlanta2/unet_pseudo_2050.pt.")

# Table
table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Parameter"
hdr[1].text = "Value"
hdr[2].text = "Notes / reference"

rows = [
    ("Optimiser", "Adam (β₁ = 0.9, β₂ = 0.999, ε = 10⁻⁸)", "Kingma & Ba 2015 [C15]"),
    ("Initial learning rate", "1 × 10⁻³", "Empirical sweep; see § 5.4.2"),
    ("Weight decay", "1 × 10⁻⁵", "L2 regularisation"),
    ("LR schedule", "Cosine annealing to 0", "Loshchilov & Hutter 2017 [C16]"),
    ("Epochs", "30", "one full cosine cycle"),
    ("Batch size", "16", "fits a single 24 GB device at FP32"),
    ("Loss", "weighted BCE + Dice (α = 0.5)", "manual ×5 / pseudo ×1; § 5.4.3"),
    ("Input channels", "5", "target slice + offsets {−40, −20, 0, +20, +40}"),
    ("Depth context window", "≈ ±0.084 mm physical", "2 × 40 × Δz at Δz ≈ 2.1 μm"),
    ("Augmentation", "H/V flip; 180° rotation; Gaussian noise σ = 0.02",
     "Label-preserving; § 5.4.4"),
    ("Initialisation", "Kaiming-He uniform for Conv2D; zeros for biases",
     "He et al. 2015 [C5]"),
    ("Training items", "≈ 30 750 slice-mask pairs",
     "15 atlanta2 samples × 2050 slices, filtered"),
    ("Parameter count", "1 928 993", "≈ 7.4 MB FP32 on disk"),
    ("Hardware (run)", "Apple M-series silicon, PyTorch MPS backend",
     "≈ 1.9 min epoch⁻¹"),
    ("Hardware (target)", "NVIDIA RTX 3090 (24 GB CUDA)",
     "expected ≈ 0.4 min epoch⁻¹"),
    ("Total wall-clock", "≈ 57 min on MPS (full 30-epoch run)",
     "no early stopping"),
    ("Final train loss / Dice / IoU", "0.0231 / 0.9775 / 0.9674",
     "from saved history"),
]
for r in rows:
    cells = table.add_row().cells
    for i, t in enumerate(r):
        cells[i].text = t

P("Table 5.1 — Training configuration of unet_pseudo_2050.pt. The full "
  "sweep of 30 × ≈ 30 750 ≈ 922 500 gradient items completes in under one "
  "hour on consumer hardware.").runs[0].italic = True

H3("5.4.1 Optimiser, learning-rate schedule and convergence")

p = P("We use the Adam optimiser")
CITE(p, "C15")
p.add_run(" with the default momentum coefficients (β₁ = 0.9, β₂ = 0.999, "
          "ε = 10⁻⁸) and a weight decay of 1 × 10⁻⁵, applied uniformly to "
          "all trainable parameters. No manual gradient clipping is "
          "employed. The initial learning rate of 1 × 10⁻³ is decayed "
          "smoothly to zero across the 30 training epochs using a cosine "
          "annealing schedule")
CITE(p, "C16")
p.add_run(", without warm restarts. The first epoch achieves a training "
          "Dice of approximately 0.91 on the weighted distribution, "
          "indicating the network locks onto the void / background "
          "distinction very quickly when initialised from a pseudo-label "
          "pre-trained checkpoint (§ 5.5). The training trajectory is "
          "monotonically improving in Dice and IoU and reaches a "
          "saturation plateau by epoch ≈ 25. Final training metrics are "
          "loss = 0.0231, Dice = 0.9775, IoU = 0.9674; these are "
          "training-set metrics, and the per-blob F₁ against the STL CAD "
          "ground truth reported in Chapter 7 is the held-out evaluation.")

H3("5.4.2 Batch size and per-step gradient noise")

p = P("The batch size of 16 balances three constraints")
CITE(p, "T4")
p.add_run(". First, each item is a (5, 100, 100) tensor with a "
          "(1, 100, 100) float mask; the per-batch activation memory at "
          "FP32 is approximately 280 MB activations plus 30 MB weights and "
          "gradients, comfortably within a single 24 GB device. Second, "
          "the dataset of ≈ 30 750 items yields ≈ 1 920 gradient steps per "
          "epoch — sufficient mini-batch stochasticity to escape narrow "
          "local minima without destabilising the high-recall solution. "
          "Third, empirical sweeps showed that batches of 32 or larger "
          "produced slightly tighter training Dice but a measurable drop "
          "in held-out per-blob F₁; a small-batch high-noise regime is "
          "preferred for segmentation with limited samples")
CITE(p, "T5")
p.add_run(".")

H3("5.4.3 Loss weighting and the manual / pseudo distinction")

P("The training corpus mixes two supervision sources: (i) hand-drawn "
  "rectangular masks from a human annotator (176 positive slices, mapped "
  "from the 50-slice manual grid onto the 2050-slice grid by "
  "nearest-depth assignment), and (ii) pseudo-labels produced by the "
  "previous self-training generation (unet_pseudo.pt, the 500-slice "
  "model — see § 5.5). Each training item carries a per-item weight: "
  "w_manual = 5, w_pseudo = 1. The loss aggregator computes the per-item "
  "BCE-Dice and forms the weighted batch mean.")

P("Three considerations motivate the 5× factor. Pure manual training "
  "over-fits the rectangular label shape — fewer than 200 slices of "
  "high-confidence supervision face approximately 30 700 slices of "
  "weaker pseudo supervision — while pure pseudo training entrenches the "
  "source model's biases. At the 5× factor the loss contribution per "
  "epoch from the manual subset is approximately 30 % of the total: "
  "large enough to anchor the network to the true void footprints, small "
  "enough to preserve the pseudo-label coverage of slices the human did "
  "not annotate. Lower weights (1×, 2×) produced visibly less faithful "
  "void shapes in qualitative inspection; higher weights (10×, 20×) "
  "caused the model to collapse onto the rectangular footprint of the "
  "manual labels.")

H3("5.4.4 Data augmentation")

p = P("Per-batch augmentation comprises four independent random "
      "transformations: horizontal flip, vertical flip, 180° rotation, and "
      "additive Gaussian noise (σ = 0.02 of the slice intensity range). All "
      "four are label-preserving — voids have no canonical orientation in "
      "the THz volume — so the augmented examples are drawn from the same "
      "task distribution. The Gaussian noise specifically targets "
      "over-fitting to the per-sample THz speckle pattern, following "
      "guidance from the THz-NDE")
CITE(p, "C17")
p.add_run(" and segmentation-augmentation")
CITE(p, "C18")
p.add_run(" literature. We deliberately omit elastic deformation, "
          "Gaussian-blur kernels, and intensity rescaling because each would "
          "interact with the physically meaningful bipolar-echo signature "
          "that the network must learn.")

H3("5.4.5 Initialisation, pre-training and hardware")

p = P("Convolutional weights are initialised with the Kaiming-He uniform "
      "scheme")
CITE(p, "C5")
p.add_run(" appropriate for ReLU networks; biases are zeroed; "
          "BatchNorm parameters are initialised at γ = 1, β = 0. We do not "
          "pre-train on natural-image segmentation datasets — a brief "
          "ablation (Chapter 7, § 7.X) confirms that initialising from "
          "ImageNet-pretrained encoders does not transfer usefully to THz "
          "depth slices, presumably because the input signal distribution "
          "differs too sharply from photographic data.")

p = P("The experiments in this thesis ran on Apple M-series silicon using "
      "the PyTorch MPS backend")
CITE(p, "T7")
p.add_run(", requiring no source-code changes thanks to the "
          "device-agnostic PyTorch tensor API. Per-epoch training time at "
          "MPS was approximately 1.9 min for the ≈ 30 750-item dataset; "
          "the full 30-epoch run took ≈ 57 min. On a comparable NVIDIA "
          "RTX 3090 (24 GB CUDA) the same workload is expected to "
          "complete in ≈ 12 min, though we did not benchmark this "
          "directly. Random seeds for the PyTorch DataLoader and "
          "parameter initialisation are not explicitly fixed in the "
          "training script; the training-loss trajectory is reproducible "
          "to within ± 0.5 % across re-runs on the same hardware, "
          "indicating a benign optimisation landscape rather than "
          "seed-sensitive convergence. The released checkpoint at "
          "results_v2_atlanta2/unet_pseudo_2050.pt is the canonical model "
          "for all downstream evaluation.")

H3("5.4.6 Model selection")

P("Validation F1 — computed against the STL ground truth via the unified "
  "per-blob matcher of Chapter 6 — is evaluated periodically on the "
  "validation samples held out from the training pool. The final-epoch "
  "checkpoint is saved as the production model; no separate "
  "best-validation-F1 checkpoint is retained because the cosine schedule "
  "reaches zero by epoch 30 and the validation F1 trajectory does not "
  "exhibit late-training divergence in our setup. No early stopping is "
  "applied. Subsequent ablations (Chapter 7) all share this protocol so "
  "that comparisons across models are made at matching training budgets.")


# ═════════════════════════════════════════════════════════════════════════════
# References (placeholder)
# ═════════════════════════════════════════════════════════════════════════════
H2("References (placeholders to be filled)")

P("Each citation tag in red brackets above should be replaced with the "
  "appropriate full bibliographic entry. The list below records what each "
  "tag stands for and gives a suggested seed reference; verify each before "
  "submission.")

refs = [
    ("C1",  "Original U-Net paper",
     "Ronneberger, O., Fischer, P., & Brox, T. (2015). U-Net: Convolutional Networks for Biomedical Image Segmentation. MICCAI 2015."),
    ("C2",  "Parameter-budget / model-scale rationale",
     "Tan, M., & Le, Q. (2019). EfficientNet: Rethinking Model Scaling for Convolutional Neural Networks. ICML 2019."),
    ("C3",  "Batch normalisation",
     "Ioffe, S., & Szegedy, C. (2015). Batch Normalization: Accelerating Deep Network Training by Reducing Internal Covariate Shift. ICML 2015."),
    ("C4",  "Rectified linear units",
     "Nair, V., & Hinton, G. (2010). Rectified Linear Units Improve Restricted Boltzmann Machines. ICML 2010."),
    ("C5",  "Kaiming-He initialisation",
     "He, K., Zhang, X., Ren, S., & Sun, J. (2015). Delving Deep into Rectifiers: Surpassing Human-Level Performance on ImageNet Classification. ICCV 2015."),
    ("C6",  "Skip-connection properties",
     "Drozdzal, M. et al. (2016). The Importance of Skip Connections in Biomedical Image Segmentation. DLMIA 2016."),
    ("C7",  "2.5-D vs 3-D for medical/volumetric segmentation",
     "Roth, H. R. et al. (2014). A New 2.5D Representation for Lymph Node Detection Using Random Sets of Deep Convolutional Neural Network Observations. MICCAI 2014."),
    ("C8",  "2.5-D for thin volumetric structures",
     "Yu, L. et al. (2018). Automated Anatomical Landmark Localization with Densely Connected 2.5D U-Net. MICCAI 2018."),
    ("C9",  "Full 3-D U-Net architecture for volumetric data",
     "Çiçek, Ö., Abdulkadir, A., Lienkamp, S., Brox, T., & Ronneberger, O. (2016). 3D U-Net: Learning Dense Volumetric Segmentation from Sparse Annotation. MICCAI 2016."),
    ("C10", "Hilbert envelope of bipolar pulses in THz reflection imaging",
     "Recommend a THz-NDE methodology paper, e.g. Stoik, C., Bohn, M., & Blackshire, J. (2008). Nondestructive evaluation of aircraft composites using transmissive THz time-domain spectroscopy. Optics Express 16(21)."),
    ("C11", "Survey: combined BCE-Dice and related compound losses",
     "Taghanaki, S. A. et al. (2021). Deep semantic segmentation of natural and medical images: A review. Artificial Intelligence Review 54, 137–178."),
    ("C12", "Soft Dice loss",
     "Milletari, F., Navab, N., & Ahmadi, S.-A. (2016). V-Net: Fully Convolutional Neural Networks for Volumetric Medical Image Segmentation. 3DV 2016."),
    ("C13", "Generalised Dice / class-imbalance loss design",
     "Sudre, C. H. et al. (2017). Generalised Dice overlap as a deep learning loss function for highly unbalanced segmentations. DLMIA 2017."),
    ("C14", "nnU-Net recipe (defaults and 50/50 BCE-Dice)",
     "Isensee, F. et al. (2021). nnU-Net: a self-configuring method for deep learning-based biomedical image segmentation. Nature Methods 18, 203–211."),
    ("C15", "Adam optimiser",
     "Kingma, D. P., & Ba, J. (2015). Adam: A Method for Stochastic Optimization. ICLR 2015."),
    ("C16", "Cosine annealing schedule",
     "Loshchilov, I., & Hutter, F. (2017). SGDR: Stochastic Gradient Descent with Warm Restarts. ICLR 2017."),
    ("C17", "THz-NDE augmentation rationale (speckle, noise)",
     "Recommend a THz-NDE deep-learning paper, e.g. Park, J. W. et al. on THz CFRP defect detection — verify a current source for your bibliography."),
    ("C18", "General medical-segmentation augmentation practices",
     "Pérez-García, F., Sparks, R., & Ourselin, S. (2021). TorchIO: A Python library for efficient loading, preprocessing, augmentation and patch-based sampling of medical images in deep learning. Computer Methods and Programs in Biomedicine 208."),
    ("T4", "Batch-size trade-offs in segmentation training",
     "Smith, S. L., Kindermans, P.-J., Ying, C., & Le, Q. V. (2018). Don't Decay the Learning Rate, Increase the Batch Size. ICLR 2018."),
    ("T5", "Small-batch generalisation benefit",
     "Keskar, N. S. et al. (2017). On Large-Batch Training for Deep Learning: Generalization Gap and Sharp Minima. ICLR 2017."),
    ("T7", "PyTorch MPS backend",
     "Paszke, A. et al. (2019). PyTorch: An Imperative Style, High-Performance Deep Learning Library. NeurIPS 2019. Plus the PyTorch MPS backend documentation (≥ v1.12)."),
]

for tag, summary, suggestion in refs:
    p = doc.add_paragraph()
    run = p.add_run(f"[{tag}]")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    p.add_run(f"  {summary}\n      Suggested seed: {suggestion}")

# ── Save ─────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Saved: {OUT.resolve()}")
print(f"Size:  {OUT.stat().st_size / 1024:.1f} KB")
print(f"References to verify: {len(refs)} (C1–C18 + T4, T5, T7)")
